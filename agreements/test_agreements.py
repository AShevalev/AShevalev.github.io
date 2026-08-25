"""Checks that the revised agreements keep KYC with LLC-FZ and do not price each leg."""

from __future__ import annotations

import re
import subprocess
import sys
from pathlib import Path

from pypdf import PdfReader

ROOT = Path(__file__).resolve().parent
OSA = ROOT / "2026-04-28_VerodusOperationalServicesAgreement.pdf"
SLA = ROOT / "2026-05-31_SoftwareLicenseAndDataLicensingAgreement.pdf"
ROYALTY = (
    "At LLC-FZ's sole discretion, and upon ninety (90) days written notice to Capital, "
    "LLC-FZ shall be entitled to charge Capital a royalty equal to five percent (5%) of "
    "gross sales (CAD), calculated and payable in quarterly installments based on the "
    "prior quarter's gross sales. The quarterly payment shall be made by Capital to "
    "LLC-FZ on or before the fifteenth (15th) day following the end of each applicable "
    "quarter period."
)


def pdf_text(path: Path) -> str:
    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def normalize(text: str) -> str:
    text = text.replace("\u2018", "'").replace("\u2019", "'")
    text = text.replace("\u201c", '"').replace("\u201d", '"')
    return re.sub(r"\s+", " ", text).strip()


def assert_contains(haystack: str, needle: str, label: str) -> None:
    if needle.lower() not in haystack.lower():
        raise AssertionError(f"{label}: missing expected text: {needle!r}")


def assert_absent(haystack: str, needle: str, label: str) -> None:
    if needle.lower() in haystack.lower():
        raise AssertionError(f"{label}: unexpectedly contains {needle!r}")


def image_count(path: Path) -> int:
    reader = PdfReader(str(path))
    last = reader.pages[-1]
    data = last["/Contents"].get_data()
    return data.count(b"FormXob")


def footer_ok(path: Path) -> None:
    reader = PdfReader(str(path))
    n = len(reader.pages)
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if f"-- {i} of {n} --" not in text.replace("\n", " "):
            # reportlab may split the footer; accept the pieces
            if not (f"-- {i} of" in text and f"{n} --" in text):
                raise AssertionError(f"{path.name} page {i}: missing footer '-- {i} of {n} --'")


def main() -> int:
    subprocess.check_call([sys.executable, str(ROOT / "generate_pdfs.py")], cwd=ROOT)
    subprocess.check_call([sys.executable, str(ROOT / "build_docx.py")], cwd=ROOT)
    osa = normalize(pdf_text(OSA))
    sla = normalize(pdf_text(SLA))

    # KYC sits with LLC-FZ, not Capital or the software owner.
    assert_contains(osa, "LLC-FZ shall, in its own name and for its own account, perform all KYC", "OSA")
    assert_contains(osa, "Capital shall not perform KYC for LLC-FZ", "OSA")
    assert_contains(osa, "is not LLC-FZ's KYC provider", "OSA")
    assert_contains(osa, "LLC-FZ is the Domain-facing principal for user onboarding", "OSA")
    assert_absent(osa, "including KYC verification of users for LLC-FZ", "OSA")
    assert_absent(osa, "KYC verification support for LLC-FZ", "OSA")
    assert_contains(sla, "LLC-FZ shall perform all KYC of Domain users itself", "SLA")
    assert_contains(sla, "1591011 B.C. LTD. is not LLC-FZ's KYC provider", "SLA")

    # No per-leg invoices; reciprocal consideration remains.
    for label, text in (("OSA", osa), ("SLA", sla)):
        assert_contains(text, "no invoices", label)
        assert_absent(text, "transfer pricing regulations", label)
        assert_absent(text, "BCICAC", label)
        assert_contains(text, "VanIAC", label)
        assert_absent(text, "Revised:", label)
        assert_absent(text, "August 25, 2026", label)
        assert_absent(text, "Software Sublicense", label)
        assert_absent(text, "merchant agreements", label)
        assert_absent(text, "prior agreements", label)
        assert_absent(text, "head licensor", label)

    assert_absent(osa, "Software License and Data Licensing Agreement", "OSA")
    assert_absent(sla, "Operational Services, Domain Usage", "SLA")

    # Original 5% royalty wording; no extra gloss inside 4.2.
    assert_contains(osa, ROYALTY, "OSA")
    assert_absent(osa, "It is not a price for any individual service or right", "OSA")
    assert_absent(osa, "only monetary exception", "OSA")
    assert_absent(osa, "NO SEPARATE PRICING", "OSA")
    assert_absent(sla, "five percent", "SLA")

    # Professional refinements requested after the first draft.
    assert_contains(osa, "independent contractors", "OSA")
    assert_contains(sla, "independent contractors", "SLA")
    assert_contains(osa, "planning metrics only", "OSA")
    assert_contains(
        osa,
        "solely to determine Capital's own payment-processing and performance-reward decisions",
        "OSA",
    )
    assert_contains(osa, "keep reasonable records of gross sales", "OSA")
    assert_contains(sla, "commercialize raw Data", "SLA")
    assert_contains(sla, "shall not grant any third party an exclusive right in Aggregated Insights", "SLA")
    assert_contains(sla, "retains all right, title, and interest in and to the Software", "SLA")
    assert_contains(sla, "may host and process Data solely as needed to operate the Software", "SLA")

    # Original signature dating: OSA undated; SLA dated May 31, 2026.
    assert_contains(sla, "Date: May 31, 2026", "SLA")
    assert_absent(osa, "Date: ____________________", "OSA")
    assert_absent(osa, "Date: May", "OSA")

    # Software grant is a licence from the owner, not a sublicence.
    assert_contains(sla, "GRANT OF LICENSE", "SLA")
    assert_contains(sla, "revocable limited license to install, access, and use the Software", "SLA")
    assert_absent(sla, "grants LLC-FZ a non-exclusive, non-transferable, revocable limited sublicense", "SLA")

    # Insights are delivered during the term, not only on the last day.
    assert_contains(sla, "with each status report under Section 4.2", "SLA")

    if image_count(OSA) != 2:
        raise AssertionError(f"OSA expected 2 signature images, got {image_count(OSA)}")
    if image_count(SLA) != 2:
        raise AssertionError(f"SLA expected 2 signature images, got {image_count(SLA)}")

    footer_ok(OSA)
    footer_ok(SLA)

    osa_docx = ROOT / "2026-04-28_VerodusOperationalServicesAgreement.docx"
    sla_docx = ROOT / "2026-05-31_SoftwareSublicenseAgreement.docx"
    if not osa_docx.exists() or not sla_docx.exists():
        raise AssertionError("expected Word copies of both original-named agreements")

    from docx import Document as WordDocument

    def word_text(path: Path) -> str:
        doc = WordDocument(str(path))
        return normalize(" ".join(p.text for p in doc.paragraphs))

    def word_drawings(path: Path) -> int:
        doc = WordDocument(str(path))
        return sum(1 for p in doc.paragraphs if p._p.xpath('.//*[local-name()="drawing"]'))

    osa_w = word_text(osa_docx)
    sla_w = word_text(sla_docx)

    # Word files keep original titles, dates, and signatures; body matches the PDFs.
    assert_contains(osa_w, "OPERATIONAL SERVICES, DOMAIN USAGE AND EVALUATION RIGHTS AGREEMENT", "OSA docx")
    assert_contains(sla_w, "SOFTWARE SUBLICENSE AND DATA LICENSING AGREEMENT", "SLA docx")
    assert_contains(osa_w, "Effective Date: April 28, 2026", "OSA docx")
    assert_contains(sla_w, "Date: May 31, 2026", "SLA docx")
    assert_contains(osa_w, ROYALTY, "OSA docx")
    assert_contains(osa_w, "Capital does not provide KYC services to LLC-FZ", "OSA docx")
    assert_contains(osa_w, "LLC-FZ is the Domain-facing principal for user onboarding", "OSA docx")
    assert_absent(osa_w, "KYC verification support for LLC-FZ", "OSA docx")
    assert_absent(osa_w, "including KYC verification of users for LLC-FZ", "OSA docx")
    assert_contains(sla_w, "GRANT OF LICENSE", "SLA docx")
    assert_contains(sla_w, "LLC-FZ shall perform all KYC of Domain users itself", "SLA docx")
    assert_absent(sla_w, "transfer pricing regulations", "SLA docx")
    assert_absent(osa_w, "BCICAC", "OSA docx")
    assert_absent(sla_w, "BCICAC", "SLA docx")
    assert_contains(osa_w, "VanIAC", "OSA docx")
    assert_contains(sla_w, "VanIAC", "SLA docx")
    assert_absent(osa_w, "Revised:", "OSA docx")
    assert_absent(sla_w, "Revised:", "SLA docx")
    assert_absent(osa_w, "merchant agreements", "OSA docx")
    assert_absent(sla_w, "prior agreements", "SLA docx")
    if word_drawings(osa_docx) != 2:
        raise AssertionError(f"OSA docx expected 2 signature drawings, got {word_drawings(osa_docx)}")
    if word_drawings(sla_docx) != 2:
        raise AssertionError(f"SLA docx expected 2 signature drawings, got {word_drawings(sla_docx)}")

    print("all agreement checks passed")
    print(f"OSA pages: {len(PdfReader(str(OSA)).pages)}")
    print(f"SLA pages: {len(PdfReader(str(SLA)).pages)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

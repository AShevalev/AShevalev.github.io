"""Patch revised body text into the original signed Word agreements.

Titles, party blocks, effective dates, and signature images stay with the
originals. Only recitals and operative clauses are updated to match the PDFs.
"""

from __future__ import annotations

import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent
ORIG_DIR = ROOT / "originals"
OUT_OSA = ROOT / "2026-04-28_VerodusOperationalServicesAgreement.docx"
OUT_SLA = ROOT / "2026-05-31_SoftwareSublicenseAgreement.docx"
ORIG_OSA = ORIG_DIR / "2026-04-28_VerodusOperationalServicesAgreement.docx"
ORIG_SLA = ORIG_DIR / "2026-05-31_SoftwareSublicenseAgreement.docx"


def has_drawing(para: Paragraph) -> bool:
    return bool(para._p.xpath('.//*[local-name()="drawing"]'))


def set_run_font(run, *, bold: bool | None = None, size: float = 12) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    sz = rPr.find(qn("w:sz"))
    if sz is None:
        sz = OxmlElement("w:sz")
        rPr.append(sz)
    sz.set(qn("w:val"), str(int(size * 2)))


def replace_text(para: Paragraph, text: str, *, bold: bool | None = None) -> None:
    if has_drawing(para):
        raise RuntimeError(f"refusing to rewrite a signed paragraph: {para.text!r}")
    p = para._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)
    run = para.add_run(text)
    set_run_font(run, bold=bold)


def insert_after(para: Paragraph, text: str, *, bold: bool | None = None) -> Paragraph:
    new_p = deepcopy(para._p)
    for child in list(new_p):
        if child.tag != qn("w:pPr"):
            new_p.remove(child)
    para._p.addnext(new_p)
    new_para = Paragraph(new_p, para._parent)
    run = new_para.add_run(text)
    set_run_font(run, bold=bold)
    return new_para


def delete_para(para: Paragraph) -> None:
    if has_drawing(para):
        raise RuntimeError("refusing to delete a signed paragraph")
    el = para._p
    el.getparent().remove(el)


def find_para(doc: Document, predicate, *, required: int = 1) -> Paragraph:
    matches = [p for p in doc.paragraphs if predicate(p.text)]
    if len(matches) != required:
        sample = [p.text[:90] for p in matches]
        raise LookupError(f"expected {required} match, got {len(matches)}: {sample}")
    return matches[0]


def starts(*prefixes: str):
    def predicate(text: str) -> bool:
        stripped = text.strip()
        return any(stripped.startswith(prefix) for prefix in prefixes)

    return predicate


def insert_block(after: Paragraph, items: list[tuple[str, bool | None]]) -> Paragraph:
    cursor = after
    for text, bold in items:
        cursor = insert_after(cursor, text, bold=bold)
    return cursor


def build_osa(src: Path, dest: Path) -> None:
    shutil.copy2(src, dest)
    doc = Document(str(dest))

    replace_text(
        find_para(doc, starts("C. LLC-FZ intends to engage Capital")),
        "C. LLC-FZ administers the Domain-facing user program and performs, in its own name and for its own account, all identity verification and know-your-customer checks of users of the Domain. Capital does not provide KYC services to LLC-FZ.",
    )
    replace_text(
        find_para(doc, starts("E. LLC-FZ grants to Capital the right to use the domain")),
        "E. LLC-FZ grants to Capital the right to use the Domain, the right to conduct platform evaluations for Capital's own payment-processing and performance-reward determinations, and the right to use LLC-FZ's affiliate software as part of traffic-driving services.",
    )
    replace_text(
        find_para(doc, starts("F. In exchange, Capital is required to drive traffic")),
        "F. In exchange, Capital is required to drive traffic to the Domain and provide the marketing services set out in this Agreement. Any commissions or profits generated from the use of the affiliate software or traffic-driving services shall be paid by Capital.",
    )

    replace_text(
        find_para(doc, starts('1.1 "Domain"')),
        '1.1 "Confidential Information" means a party\'s non-public information disclosed in connection with this Agreement, including business, technical, user-status, and payment-rail information. Confidential Information does not include information that is public other than through a breach of this Agreement, already lawfully known to the receiving party, independently developed without use of the disclosing party\'s information, or lawfully received from a third party without a duty of confidentiality.',
    )
    insert_block(
        find_para(doc, starts('1.1 "Confidential Information"')),
        [('1.2 "Domain" means the domain name verodus.com.', None)],
    )
    replace_text(
        find_para(doc, starts('1.2 "Services"')),
        '1.3 "KYC" means identity verification, know-your-customer, anti-money-laundering, and counter-terrorist-financing onboarding checks of Domain users.',
    )
    insert_block(
        find_para(doc, starts('1.3 "KYC"')),
        [
            (
                '1.4 "Services" means the marketing and traffic-driving services provided by Capital to LLC-FZ as described in Section 2.1.',
                None,
            )
        ],
    )
    replace_text(
        find_para(doc, starts('1.3 "Effective Date"', '1.5 "Effective Date"')),
        '1.5 "Effective Date" means the date first written above.',
    )

    replace_text(
        find_para(doc, starts("(a) Administrative and operational support")),
        "(a) Marketing and traffic-driving services directed exclusively to the Domain, including the use of LLC-FZ's affiliate software.",
    )
    delete_para(
        find_para(
            doc,
            starts("(b) Marketing and traffic-driving services directed exclusively to the Domain"),
        )
    )
    replace_text(
        find_para(doc, lambda t: t.strip() == "(b) Platform evaluations."),
        "(b) Platform evaluations, solely to determine Capital's own payment-processing and performance-reward decisions.",
    )
    replace_text(
        find_para(doc, starts("2.3 Capital shall perform all payment processing")),
        "2.3 Capital shall perform all payment processing on its own rails and platform evaluations on its own behalf and shall have no access to or involvement in the software or merchant accounts of LLC-FZ for LLC-FZ's isolated payment rails.",
    )
    p24 = find_para(doc, starts("2.4 LLC-FZ shall independently establish"))
    replace_text(
        p24,
        "2.4 LLC-FZ shall independently establish and operate its own isolated merchant accounts for its payment rails (deposits and same-method payouts only). There shall be no crossover or commingling of funds between Capital's payment rails and LLC-FZ's payment rails. Capital has no operational control over LLC-FZ's merchant accounts.",
    )
    insert_block(
        p24,
        [
            (
                "2.5 LLC-FZ is the Domain-facing principal for user onboarding. LLC-FZ shall, in its own name and for its own account, perform all KYC of Domain users.",
                None,
            ),
            (
                "(a) Capital shall not perform KYC for LLC-FZ, shall not collect identity documents or KYC files on LLC-FZ's behalf, and is not LLC-FZ's KYC provider, outsourced compliance function, or KYC processor.",
                None,
            ),
            (
                "(b) LLC-FZ may give Capital only the minimum user-status information reasonably required for Capital's independent payment-processing and performance-reward activities (for example, verified, unverified, or restricted). LLC-FZ is not required to give Capital underlying KYC documents unless Capital's payment processors or applicable law require it, in which case LLC-FZ shall provide only what is required.",
                None,
            ),
            (
                "(c) Capital may collect and process payment-instrument, fraud-prevention, sanctions, and settlement data on its own payment rails as required by its processors and applicable law. That activity is Capital's own payment-rail compliance. It is not KYC performed for LLC-FZ.",
                None,
            ),
        ],
    )

    replace_text(
        find_para(doc, starts("3.1 LLC-FZ hereby grants Capital")),
        "3.1 LLC-FZ hereby grants Capital a non-exclusive, non-transferable, revocable license to use the Domain, conduct platform evaluations, and use LLC-FZ's affiliate software solely for marketing, traffic-driving, and Capital's independent activities under Section 2.2 during the Term.",
    )
    insert_block(
        find_para(doc, starts("Any commissions or profits generated through the affiliate software")),
        [
            (
                "3.4 Capital shall not hold itself out as LLC-FZ, except to the extent reasonably necessary to use the Domain under this Section 3. Capital shall not bind LLC-FZ to any obligation.",
                None,
            )
        ],
    )
    insert_block(
        find_para(doc, starts("4.2 At LLC-FZ")),
        [
            (
                "4.3 If LLC-FZ delivers notice under Section 4.2, Capital shall keep reasonable records of gross sales sufficient to calculate the royalty and, on reasonable written request, make those records available to LLC-FZ while the royalty is payable.",
                None,
            )
        ],
    )
    insert_block(
        find_para(doc, starts("5.2 Targets shall be reviewed")),
        [
            (
                "5.3 The targets are planning metrics only. Failure to agree on or meet a target is not a material breach and does not change the rights or consideration under this Agreement.",
                None,
            )
        ],
    )

    replace_text(
        find_para(doc, starts("(c) It will comply with all applicable laws, including anti-money-laundering")),
        "(c) It will comply with all applicable laws that apply to its own activities, including anti-money-laundering and payment-processing regulations.",
    )
    insert_block(
        find_para(
            doc,
            starts("(c) It will comply with all applicable laws that apply to its own activities, including anti-money-laundering"),
        ),
        [
            ("6.2 LLC-FZ represents that it, and not Capital, is responsible for KYC of Domain users.", None),
            (
                "6.3 Capital represents that it is responsible for legal and processor obligations arising from merchant accounts held in Capital's name.",
                None,
            ),
        ],
    )

    replace_text(find_para(doc, lambda t: t.strip() == "7. CONFIDENTIALITY"), "7. CONFIDENTIALITY AND SECURITY", bold=True)
    replace_text(
        find_para(doc, starts("7.1 Each party shall maintain the confidentiality")),
        "7.1 Each party shall keep the other party's Confidential Information strictly confidential and use it solely for the purposes of this Agreement. KYC documents and identity-verification files of Domain users are LLC-FZ's Confidential Information. Capital shall not use or retain them except as Section 2.5(b) or applicable law requires.",
    )
    insert_block(
        find_para(doc, starts("7.1 Each party shall keep the other party's Confidential Information")),
        [
            (
                "7.2 A party may disclose Confidential Information to its employees, contractors, professional advisors, and processors who have a need to know and are bound to confidentiality no less protective than this Section 7, or if required by law, a regulator, or a payment processor, in which case the disclosing party shall, if legally permitted, give the other party prompt notice.",
                None,
            ),
            (
                "7.3 Each party shall apply reasonable technical and organizational measures to protect Confidential Information and any user-status or payment data in its possession under this Agreement.",
                None,
            ),
            ("7.4 The obligations in this Section 7 survive termination.", None),
        ],
    )

    replace_text(
        find_para(doc, starts("8.3 Upon termination, the Domain and Evaluation Rights")),
        "8.3 Upon termination, the Domain and Evaluation Rights license granted herein shall terminate immediately. Capital shall stop using the Domain and the affiliate software, and shall return or securely delete KYC documents (if any) in its possession, except as applicable law requires Capital to retain. Accrued payment obligations under Section 4.2, if any, remain payable.",
    )
    replace_text(
        find_para(doc, starts("9.1 Each party shall indemnify the other against claims")),
        "9.1 Each party shall indemnify the other against claims, losses, or damages arising from its own breach or violation of law, including, in LLC-FZ's case, claims arising from LLC-FZ's KYC of Domain users, and in Capital's case, claims arising from Capital's merchant accounts and payment rails.",
    )
    insert_block(
        find_para(doc, starts("10.1 Neither party shall be liable for indirect, consequential, or punitive damages.")),
        [
            (
                "10.2 Section 10.1 does not limit a party's liability for fraud, wilful misconduct, or its indemnification obligations under Section 9.",
                None,
            )
        ],
    )
    replace_text(
        find_para(doc, lambda t: t.startswith("11.2") and "BCICAC" in t),
        "11.2 Any disputes arising out of or in connection with this Agreement shall be finally settled by arbitration administered by the Vancouver International Arbitration Centre (VanIAC) in Vancouver, British Columbia, before a single arbitrator, in accordance with the VanIAC International Commercial Arbitration Rules. The language of arbitration shall be English.",
    )

    replace_text(
        find_para(doc, starts("12.1 This Agreement is bilateral")),
        "12.1 This Agreement is bilateral.",
    )
    replace_text(
        find_para(doc, starts("12.2 Capital has no access or rights to the software")),
        "12.2 Capital has no access or rights to the software of LLC-FZ (other than the limited license to use the affiliate software for traffic-driving purposes), no operational involvement in LLC-FZ's isolated merchant accounts, and no KYC role for LLC-FZ.",
    )
    insert_block(
        find_para(doc, starts("12.2 Capital has no access or rights to the software")),
        [
            (
                "12.3 The parties are independent contractors. This Agreement does not create a partnership, joint venture, fiduciary relationship, or agency, and neither party may bind the other.",
                None,
            ),
            (
                "12.4 Notices under this Agreement must be in writing and delivered by registered mail to the recipient's registered and records office. A notice is effective on delivery or, if delivery is refused, on the date of attempted delivery. Notice under Section 4.2 must be given in the same manner.",
                None,
            ),
            (
                "12.5 Neither party may assign this Agreement without the other party's prior written consent, except to an affiliate or to a successor in connection with a merger or a sale of substantially all of its assets. Any other purported assignment is void.",
                None,
            ),
            (
                "12.6 Sections 4.2 and 4.3 (to the extent a royalty has become payable), 7, 8.3, 9, 10, 11, and 12 survive termination.",
                None,
            ),
            (
                "12.7 Failure to enforce a provision is not a waiver. A waiver must be in writing. If any provision is held invalid, the remainder of this Agreement shall continue in full force and effect. Headings are for convenience only.",
                None,
            ),
            ("12.8 This Agreement does not confer rights on any person who is not a party.", None),
        ],
    )
    replace_text(
        find_para(doc, starts("12.3 This Agreement constitutes the entire understanding")),
        "12.9 This Agreement constitutes the entire understanding between the parties and supersedes all prior discussions.",
    )
    replace_text(
        find_para(doc, starts("12.4 Any amendments must be in writing")),
        "12.10 Any amendments must be in writing and signed by both parties.",
    )
    delete_para(find_para(doc, starts("12.5 If any provision is held invalid")))
    replace_text(
        find_para(doc, starts("12.6 This Agreement may be executed in counterparts")),
        "12.11 This Agreement may be executed in counterparts, including electronically.",
    )

    doc.save(str(dest))


def build_sla(src: Path, dest: Path) -> None:
    shutil.copy2(src, dest)
    doc = Document(str(dest))

    replace_text(
        find_para(doc, starts("C. LLC-FZ is the registered owner of the domain name verodus.com.")),
        "C. LLC-FZ is the registered owner of the domain name verodus.com. LLC-FZ performs, in its own name and for its own account, all identity verification and know-your-customer checks of users of that domain.",
    )
    replace_text(
        find_para(doc, starts("D. The parties desire to enter into a bilateral arrangement")),
        "D. The parties desire to enter into a bilateral arrangement whereby 1591011 B.C. LTD. grants LLC-FZ a license to use the Software for data collection and analysis (including LLC-FZ's own KYC and classification activities), and LLC-FZ in return grants 1591011 B.C. LTD. a license to the Aggregated Insights generated through that use.",
    )
    replace_text(
        find_para(doc, starts('1.1 "Aggregated Insights"')),
        '1.1 "Aggregated Insights" means anonymized, amalgamated data and analytical outputs derived from the application of the Software, excluding any raw client-identifying information, KYC documents, and identity-verification files.',
    )
    insert_block(
        find_para(doc, starts('1.1 "Aggregated Insights"')),
        [
            (
                '1.2 "Confidential Information" means a party\'s non-public information disclosed in connection with this Agreement, including the Software, Data, KYC records, and Aggregated Insights before they are made public. Confidential Information does not include information that is public other than through a breach of this Agreement, already lawfully known to the receiving party, independently developed without use of the disclosing party\'s information, or lawfully received from a third party without a duty of confidentiality.',
                None,
            )
        ],
    )
    replace_text(
        find_para(doc, starts('1.2 "Data"')),
        '1.3 "Data" means all trading performance metrics, client transaction data, and any other information collected or generated through the Software.',
    )
    insert_block(
        find_para(doc, starts('1.3 "Data"')),
        [('1.4 "Domain" means the domain name verodus.com.', None)],
    )
    replace_text(
        find_para(doc, starts('1.3 "Effective Date"', '1.5 "Effective Date"')),
        '1.5 "Effective Date" means the date first written above.',
    )
    insert_block(
        find_para(doc, starts('1.5 "Effective Date"')),
        [
            (
                '1.6 "KYC" means identity verification, know-your-customer, anti-money-laundering, and counter-terrorist-financing onboarding checks of Domain users.',
                None,
            )
        ],
    )
    replace_text(
        find_para(doc, starts('1.4 "Software"', '1.7 "Software"')),
        '1.7 "Software" means the pre-existing software, CRM systems, and related technology owned and operated by 1591011 B.C. LTD.',
    )

    replace_text(find_para(doc, lambda t: t.strip() == "2. GRANT OF SUBLICENSE"), "2. GRANT OF LICENSE", bold=True)
    replace_text(
        find_para(doc, starts("2.1 1591011 B.C. LTD. hereby grants LLC-FZ")),
        "2.1 1591011 B.C. LTD. hereby grants LLC-FZ a non-exclusive, non-transferable, revocable limited license to install, access, and use the Software in connection with LLC-FZ's Meydan-authorized business activities, including data collection, classification, analysis, and LLC-FZ's own KYC of Domain users.",
    )
    replace_text(
        find_para(doc, starts("2.2 The sublicense does not include", "2.2 The license does not include")),
        "2.2 The license does not include any right to modify, reverse-engineer, decompile, copy (except for internal backup), or sublicense the Software. LLC-FZ shall not sublicense the Software to any third party under any circumstance.",
    )
    insert_block(
        find_para(doc, starts("2.2 The license does not include")),
        [
            (
                "2.3 1591011 B.C. LTD. retains all right, title, and interest in and to the Software. LLC-FZ acquires only the license granted in this Section 2. Any improvement, configuration, or derivative of the Software remains the exclusive property of 1591011 B.C. LTD.",
                None,
            )
        ],
    )

    replace_text(
        find_para(doc, starts("3.1 LLC-FZ shall collect, classify, analyze")),
        "3.1 LLC-FZ shall collect, classify, analyze, and aggregate Data through the Software, including trading performance metrics across supported instruments. LLC-FZ shall perform all KYC of Domain users itself. LLC-FZ shall not be required to share KYC documents or identity-verification files with 1591011 B.C. LTD.",
    )
    replace_text(
        find_para(doc, starts("3.2 All raw Data")),
        "3.2 All raw Data, client-identifying information, and KYC records derived from or stored in the Software shall be the sole and exclusive property of LLC-FZ.",
    )
    replace_text(
        find_para(doc, starts("3.3 LLC-FZ hereby grants 1591011 B.C. LTD. a perpetual")),
        "3.3 LLC-FZ hereby grants 1591011 B.C. LTD. a perpetual, irrevocable, royalty-free, worldwide, non-exclusive license to use, reproduce, modify, adapt, resell, sublicense to third parties, commercialize and otherwise exploit the Aggregated Insights in any manner whatsoever. That license does not include raw Data, client-identifying information, or KYC records.",
    )
    replace_text(
        find_para(doc, starts("3.4 LLC-FZ is expressly permitted to sell")),
        "3.4 LLC-FZ is expressly permitted to sell, license, or otherwise commercialize raw Data to any third party at its sole discretion, subject to applicable law, provided it does not grant any third party a license to the Software itself. LLC-FZ may use Aggregated Insights in its own business. Section 3.3 is the grant under which 1591011 B.C. LTD. may commercialize Aggregated Insights. LLC-FZ shall not grant any third party an exclusive right in Aggregated Insights that conflicts with Section 3.3.",
    )
    insert_block(
        find_para(doc, starts("3.4 LLC-FZ is expressly permitted")),
        [
            (
                "3.5 1591011 B.C. LTD. may host and process Data solely as needed to operate the Software for LLC-FZ and to generate Aggregated Insights under Section 3.3. 1591011 B.C. LTD. shall not sell personal Data, shall not use personal Data for its own marketing, and shall not disclose personal Data except to subprocessors bound to confidentiality who need access to operate the Software, or as required by law.",
                None,
            ),
            (
                "3.6 Each party shall apply reasonable technical and organizational measures to protect Data, KYC records, and the Software. 1591011 B.C. LTD. shall not include KYC records or raw client-identifying information in Aggregated Insights.",
                None,
            ),
        ],
    )

    replace_text(
        find_para(doc, starts("4.1 LLC-FZ shall deliver the complete set of Aggregated Insights")),
        "4.1 LLC-FZ shall deliver then-available Aggregated Insights to 1591011 B.C. LTD. with each status report under Section 4.2, and shall deliver the complete set of Aggregated Insights generated during the Term no later than the last day of the Term.",
    )
    replace_text(
        find_para(doc, starts("5.1 The parties acknowledge and agree that the only consideration")),
        "5.1 The parties acknowledge and agree that the only consideration for the license and the data license is the reciprocal grant of rights set forth in Sections 2 and 3.",
    )
    replace_text(
        find_para(doc, starts("5.3 The parties further acknowledge")),
        "5.3 The parties agree that the reciprocal licenses and rights granted under this Agreement are sufficient consideration for this Agreement.",
    )

    replace_text(
        find_para(doc, starts("(c) It will comply with all applicable laws, including data protection")),
        "(c) It will comply with all applicable laws that apply to its own activities, including data protection and export-control laws.",
    )
    insert_block(
        find_para(
            doc,
            starts("(c) It will comply with all applicable laws that apply to its own activities, including data protection"),
        ),
        [
            (
                "6.2 LLC-FZ represents that it, and not 1591011 B.C. LTD., is responsible for KYC of Domain users and for privacy and anti-money-laundering laws applicable to that activity. 1591011 B.C. LTD. is not LLC-FZ's KYC provider.",
                None,
            ),
            (
                "6.3 Except as expressly set out in this Agreement, 1591011 B.C. LTD. makes no representations or warranties regarding the Software, whether express, implied, or statutory.",
                None,
            ),
        ],
    )
    replace_text(
        find_para(doc, starts("7.1 Each party shall keep the other party's non-public information")),
        "7.1 Each party shall keep the other party's Confidential Information strictly confidential and use it only for the purposes of this Agreement. KYC records and client-identifying information are LLC-FZ's Confidential Information. 1591011 B.C. LTD. shall not use them except as technically required to operate the Software for LLC-FZ, and shall not include them in Aggregated Insights.",
    )
    insert_block(
        find_para(doc, starts("7.1 Each party shall keep the other party's Confidential Information")),
        [
            (
                "7.2 A party may disclose Confidential Information to its employees, contractors, professional advisors, and subprocessors who have a need to know and are bound to confidentiality no less protective than this Section 7, or if required by law or a regulator, in which case the disclosing party shall, if legally permitted, give the other party prompt notice.",
                None,
            ),
            ("7.3 The obligations in this Section 7 survive termination.", None),
        ],
    )
    replace_text(
        find_para(doc, starts("8.3 Upon termination, the sublicense granted under Section 2")),
        "8.3 Upon termination, the license granted under Section 2 ends immediately, but the perpetual license to Aggregated Insights granted under Section 3.3 survives. 1591011 B.C. LTD. shall not retain KYC records or raw client-identifying information after termination except as required by law, or as residual Aggregated Insights already delivered in anonymized form.",
    )
    replace_text(
        find_para(doc, starts("9.1 Each party shall indemnify, defend, and hold harmless")),
        "9.1 Each party shall indemnify, defend, and hold harmless the other party against any claims, losses, or damages arising from its own breach of this Agreement or violation of applicable law, including, in LLC-FZ's case, claims arising from LLC-FZ's KYC of Domain users or LLC-FZ's handling of Data.",
    )
    insert_block(
        find_para(doc, starts("10.1 Neither party shall be liable for indirect, consequential, punitive")),
        [
            (
                "10.2 Section 10.1 does not limit a party's liability for fraud, wilful misconduct, or its indemnification obligations under Section 9.",
                None,
            )
        ],
    )
    replace_text(
        find_para(doc, lambda t: t.startswith("11.2") and "BCICAC" in t),
        "11.2 Any disputes arising out of or in connection with this Agreement shall be finally settled by arbitration administered by the Vancouver International Arbitration Centre (VanIAC) in Vancouver, British Columbia, before a single arbitrator, in accordance with the VanIAC International Commercial Arbitration Rules. The language of arbitration shall be English.",
    )

    replace_text(
        find_para(doc, starts("12.1 This Agreement is bilateral")),
        "12.1 This Agreement is bilateral.",
    )
    insert_block(
        find_para(doc, starts("12.1 This Agreement is bilateral.")),
        [
            (
                "12.2 The parties are independent contractors. This Agreement does not create a partnership, joint venture, fiduciary relationship, or agency, and neither party may bind the other.",
                None,
            ),
            (
                "12.3 Notices under this Agreement must be in writing and delivered by registered mail to the recipient's registered and records office. A notice is effective on delivery or, if delivery is refused, on the date of attempted delivery.",
                None,
            ),
            (
                "12.4 Neither party may assign this Agreement without the other party's prior written consent, except to an affiliate or to a successor in connection with a merger or a sale of substantially all of its assets. Any other purported assignment is void.",
                None,
            ),
            (
                "12.5 Sections 3.3, 3.4, 3.5, 3.6, 7, 8.3, 9, 10, 11, and 12 survive termination.",
                None,
            ),
            (
                "12.6 Failure to enforce a provision is not a waiver. A waiver must be in writing. If any provision is held invalid, the remainder of this Agreement shall continue to be binding. Headings are for convenience only.",
                None,
            ),
            ("12.7 This Agreement does not confer rights on any person who is not a party.", None),
        ],
    )
    replace_text(
        find_para(doc, starts("12.2 This Agreement constitutes the entire understanding")),
        "12.8 This Agreement constitutes the entire understanding and supersedes all prior discussions.",
    )
    replace_text(
        find_para(doc, starts("12.3 No amendment is valid")),
        "12.9 No amendment is valid unless in writing and signed by both parties.",
    )
    delete_para(find_para(doc, starts("12.4 If any provision is held invalid")))
    replace_text(
        find_para(doc, starts("12.5 This Agreement may be executed in counterparts")),
        "12.10 This Agreement may be executed in counterparts, including electronically.",
    )

    doc.save(str(dest))


def main() -> None:
    build_osa(ORIG_OSA, OUT_OSA)
    build_sla(ORIG_SLA, OUT_SLA)
    print(f"wrote {OUT_OSA}")
    print(f"wrote {OUT_SLA}")


if __name__ == "__main__":
    main()
